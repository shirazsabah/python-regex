# -*- coding: utf-8 -*-
"""
Python script to create Word documents with random words (including prespecified keywords)
"""

# ## Install python-docx if not already installed
# pip install python-docx

# ## This script assumes that the documents are on the Desktop in a folder called 'PNI letters'

import random
import string
from docx import Document
from datetime import datetime
import os

# List of keywords
keywords = ["total hip replacement", "hip replacement", "THR", "sciatic nerve"]

# Function to generate a random string of uppercase letters
def random_string(length):
    letters = string.ascii_uppercase
    return ''.join(random.choice(letters) for _ in range(length))

# Set the working directory to the desktop on Windows
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop/PNI letters")
os.chdir(desktop_path)

# Generate 10000 Word documents
for doc_num in range(1, 10001):
    # Generate an 8-digit random number
    random_number = ''.join(random.choices(string.digits, k=8))
    
    # Generate a random firstname and surname
    firstname = random_string(5)
    surname = random_string(8)
    
    # Generate a random date
    year = random.randint(2010, 2023)  # Random year between 2010 and 2023
    month = random.randint(1, 12)  # Random month between 1 and 12
    day = random.randint(1, 28)  # Random day between 1 and 28
    random_date = datetime(year, month, day).strftime("%Y-%m-%d")
    
    # Create the document name
    document_name = f"{random_number}-{firstname} {surname}-{random_date}"
    
    # Create a new document
    document = Document()
    
    # Generate 1000 random words
    for _ in range(1000):
        # Generate a random word
        word = ''.join(random.choice('abcdefghijklmnopqrstuvwxyz') for _ in range(random.randint(3, 10)))
        # Add the word to the document
        document.add_paragraph(word)
    
    # Randomly select documents to contain keywords
    if random.random() < 0.2:  # 20% chance of containing keywords
        # Choose two random keywords
        selected_keywords = random.sample(keywords, 2)
        # Add the keywords to the document
        for keyword in selected_keywords:
            document.add_paragraph(keyword)
    
        # Save the document with the generated name
    document.save(f'{document_name}.docx')
